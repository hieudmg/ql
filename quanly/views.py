# coding=utf-8

from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string
from .models import *
from .forms import *
from datetime import *
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import *
from django.core.files.storage import FileSystemStorage
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm


def thongtin_list(request):
    thongtin = ThongTin.objects.all()
    return render(request, 'benhnhan/thongtin_list.html', {'thongtin': thongtin})


def test(request):
    return render(request, 'benhnhan/test.html')


def thongtin_add(request):
    data = dict()

    if request.method == 'POST':
        form = FormTT(request.POST)
        if form.is_valid():
            form.save()
            tt = ThongTin.objects.latest('id')
            tr = Trung(tt=tt)
            tr.save()
            tp = TruPhoi(tt=tt)
            tp.save()
            p = Phoi(tt=tt)
            p.save()
            td = TinhDichDoNgayCH(tt=tt)
            td.save()
            ch = ChocHut(tt=tt)
            ch.save()
            cp = ChuyenPhoi(tt=tt)
            cp.save()
            dp = DongPhoi(tt=tt)
            dp.save()
            data['form_is_valid'] = True
            thongtin = ThongTin.objects.all()
            data['html_thongtin_preview'] = render_to_string('benhnhan/includes/thongtin_preview.html', {
                'thongtin': thongtin, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        form = FormTT()

    context = {'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/thongtin_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def trung_add(request):
    data = dict()
    tt = ThongTin.objects.latest('id')
    tr = Trung(tt=tt)

    if request.method == 'POST':
        formtr = FormTR(request.POST, instance=tr)
        if formtr.is_valid():
            formtr.save()
            data['formtr_is_valid'] = True
        else:
            data['formtr_is_valid'] = False
    else:
        formtr = FormTR(instance=tr)

    context = {'formtr': formtr}
    data['html_formtr'] = render_to_string('benhnhan/includes/trung_add.html',
                                           context,
                                           request=request
                                           )
    return JsonResponse(data)


def thongtin_edit(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)

    if request.method == 'POST':
        form = FormTT(request.POST, instance=thongtin)
        formtr = FormTR(request.POST, instance=thongtin.trung)
        formp = FormP(request.POST, instance=thongtin.phoi)
        formtp = FormTP(request.POST, instance=thongtin.truphoi)
        formtd = FormTD(request.POST, instance=thongtin.tinhdichdongaych)
        if formtd.is_valid() and formtr.is_valid() and formp.is_valid() and formtp.is_valid():
            # form.save()
            formtr.save()
            formp.save()
            formtp.save()
            formtd.save()
            data['form_is_valid'] = True
            thongtin = ThongTin.objects.all()
            data['html_thongtin_preview'] = render_to_string('benhnhan/includes/thongtin_preview.html', {
                'thongtin': thongtin, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        form = FormTT(instance=thongtin)
        formtr = FormTR(instance=thongtin.trung)
        formp = FormP(instance=thongtin.phoi)
        formtp = FormTP(instance=thongtin.truphoi)
        formtd = FormTD(instance=thongtin.tinhdichdongaych)
    context = {'form': form, 'formtr': formtr, 'formp': formp, 'formtp': formtp, 'formtd': formtd}
    data['html_form'] = render_to_string('benhnhan/includes/thongtin_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def thongtin_info(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    if request.method == 'POST':
        form = FormTT(request.POST, instance=thongtin)
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True
            thongtin = ThongTin.objects.all()
            data['html_thongtin_preview'] = render_to_string('benhnhan/includes/thongtin_preview.html', {
                'thongtin': thongtin, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        form = FormTT(instance=thongtin)

    context = {'form': form, }
    data['html_form'] = render_to_string('benhnhan/includes/thongtin_info.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def thongtin_del(request, pk):
    tt_del = get_object_or_404(ThongTin, pk=pk)
    data = dict()
    if request.method == 'POST':
        tt_del.delete()
        data['form_is_valid'] = True  # This is just to play along with the existing code
        thongtin = ThongTin.objects.all()
        data['html_thongtin_preview'] = render_to_string('benhnhan/includes/thongtin_preview.html', {
            'thongtin': thongtin, 'user': request.user
        })
    else:
        context = {'tt_del': tt_del}
        data['html_form'] = render_to_string('benhnhan/includes/thongtin_del.html', context, request=request)
    return JsonResponse(data)


def bacsi_list(request):
    bacsi = BacSi.objects.all()
    return render(request, 'bacsi/bacsi_list.html', {'bacsi': bacsi})


def bacsi_add(request):
    data = dict()

    if request.method == 'POST':
        formbs = FormBS(request.POST)
        if formbs.is_valid():
            formbs.save()
            bacsi = BacSi.objects.all()
            data['form_is_valid'] = True
            data['html_bacsi_preview'] = render_to_string('bacsi/includes/bacsi_preview.html', {
                'bacsi': bacsi, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        formbs = FormBS()

    context = {'formbs': formbs}
    data['html_form'] = render_to_string('bacsi/includes/bacsi_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def bacsi_edit(request, pk):
    data = dict()
    bacsi = get_object_or_404(BacSi, pk=pk)
    if request.method == 'POST':
        formbs = FormBS(request.POST, instance=bacsi)
        if formbs.is_valid():
            formbs.save()
            data['form_is_valid'] = True
            bacsi = BacSi.objects.all()
            data['html_bacsi_preview'] = render_to_string('bacsi/includes/bacsi_preview.html', {
                'bacsi': bacsi, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        formbs = FormBS(instance=bacsi)

    context = {'formbs': formbs, }
    data['html_form'] = render_to_string('bacsi/includes/bacsi_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def bacsi_del(request, pk):
    data = dict()
    bacsi = get_object_or_404(BacSi, pk=pk)
    if request.method == 'POST':
        bacsi.delete()
        data['form_is_valid'] = True
        bacsi = BacSi.objects.all()
        data['html_bacsi_preview'] = render_to_string('bacsi/includes/bacsi_preview.html', {
            'bacsi': bacsi, 'user': request.user
        })
    else:
        context = {'bacsi': bacsi}
        data['html_form'] = render_to_string('bacsi/includes/bacsi_del.html', context, request=request)
    return JsonResponse(data)


def kythuatvien_list(request):
    kythuatvien = KyThuatVien.objects.all()
    return render(request, 'kythuatvien/kythuatvien_list.html', {'kythuatvien': kythuatvien})


def kythuatvien_add(request):
    data = dict()

    if request.method == 'POST':
        formktv = FormKTV(request.POST)
        if formktv.is_valid():
            formktv.save()
            kythuatvien = KyThuatVien.objects.all()
            data['form_is_valid'] = True
            data['html_kythuatvien_preview'] = render_to_string('kythuatvien/includes/kythuatvien_preview.html', {
                'kythuatvien': kythuatvien, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        formktv = FormKTV()

    context = {'formktv': formktv}
    data['html_form'] = render_to_string('kythuatvien/includes/kythuatvien_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def kythuatvien_edit(request, pk):
    data = dict()
    kythuatvien = get_object_or_404(KyThuatVien, pk=pk)
    if request.method == 'POST':
        formktv = FormKTV(request.POST, instance=kythuatvien)
        if formktv.is_valid():
            formktv.save()
            data['form_is_valid'] = True
            kythuatvien = KyThuatVien.objects.all()
            data['html_kythuatvien_preview'] = render_to_string('kythuatvien/includes/kythuatvien_preview.html', {
                'kythuatvien': kythuatvien, 'user': request.user
            })
        else:
            data['form_is_valid'] = False
    else:
        formktv = FormKTV(instance=kythuatvien)

    context = {'formktv': formktv, }
    data['html_form'] = render_to_string('kythuatvien/includes/kythuatvien_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def kythuatvien_del(request, pk):
    data = dict()
    kythuatvien = get_object_or_404(KyThuatVien, pk=pk)
    if request.method == 'POST':
        kythuatvien.delete()
        data['form_is_valid'] = True
        kythuatvien = KyThuatVien.objects.all()
        data['html_kythuatvien_preview'] = render_to_string('kythuatvien/includes/kythuatvien_preview.html', {
            'kythuatvien': kythuatvien, 'user': request.user
        })
    else:
        context = {'kythuatvien': kythuatvien}
        data['html_form'] = render_to_string('kythuatvien/includes/kythuatvien_del.html', context, request=request)
    return JsonResponse(data)


def chochut_stt():
    chochut = list(ChocHut.objects.filter(added=True).order_by('gioCH'))
    chochut[0].stt = 1
    chochut[0].save()
    for i in range(0, chochut.__len__() - 1, 1):
        if chochut[i].gioCH.date() == chochut[i + 1].gioCH.date():
            chochut[i + 1].stt = chochut[i].stt + 1
        else:
            chochut[i + 1].stt = 1
        chochut[i + 1].save()


def chochut_list(request):
    chochut_stt()
    data = dict()
    try:
        ngay = datetime.strptime(str(request.GET.get('d')), '%d/%m/%Y')
        data['chonngay'] = True
        chochut = ChocHut.objects.filter(added=True).filter(gioCH__date=ngay)
    except ValueError:
        data['chonngay'] = False
        chochut = ChocHut.objects.filter(added=True)
    return render(request, 'benhnhan/chochut_list.html', {'chochut': chochut, 'data': data})


def chochut_add(request, pk):
    data = dict()

    thongtin = get_object_or_404(ThongTin, pk=pk)
    chochut = thongtin.chochut
    data['added'] = True
    data['maso'] = thongtin.maSo
    if not chochut.added:
        chochut.HCG = datetime.now
        chochut.gioCH = datetime.now
        data['added'] = False

    if request.method == 'POST':
        formch = FormCH(request.POST, instance=chochut)
        form = FormTT(request.POST, instance=thongtin)
        if formch.is_valid():
            formch.save()
            chochut_stt()
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formch = FormCH(instance=chochut)
        form = FormTT(instance=thongtin)

    context = {'formch': formch, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/chochut_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def chochut_edit(request, pk):
    data = dict()
    ngay = ''
    try:
        ngay = datetime.strptime(str(request.GET.get('d')), '%d/%m/%Y')
        data['chonngay'] = True
    except ValueError:
        data['chonngay'] = False

    thongtin = get_object_or_404(ThongTin, pk=pk)
    chochut = thongtin.chochut
    if request.method == 'POST':
        formch = FormCH(request.POST, instance=chochut)
        form = FormTT(instance=thongtin)
        chochut_stt()
        if formch.is_valid():
            formch.save()
            if data['chonngay']:
                chochut = ChocHut.objects.filter(added=True).filter(gioCH__date=ngay)
            else:
                chochut = ChocHut.objects.filter(added=True)
            data['html_chochut_preview'] = render_to_string('benhnhan/includes/chochut_preview.html', {
                'chochut': chochut, 'user': request.user, 'data': data
            })
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formch = FormCH(instance=chochut)
        form = FormTT(instance=thongtin)

    context = {'formch': formch, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/chochut_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def chochut_del(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    chochut = thongtin.chochut
    ngay = ''
    try:
        ngay = datetime.strptime(str(request.GET.get('d')), '%d/%m/%Y')
        data['chonngay'] = True
    except ValueError:
        data['chonngay'] = False

    if request.method == 'POST':
        chochut_stt()
        chochut.added = False
        chochut.save()
        data['form_is_valid'] = True
        if data['chonngay']:
            chochut = ChocHut.objects.filter(added=True).filter(gioCH__date=ngay)
        else:
            chochut = ChocHut.objects.filter(added=True)
        data['html_chochut_preview'] = render_to_string('benhnhan/includes/chochut_preview.html', {
            'chochut': chochut, 'user': request.user, 'data': data
        })
    else:
        context = {'chochut': chochut}
        data['html_form'] = render_to_string('benhnhan/includes/chochut_del.html', context, request=request)
    return JsonResponse(data)


def chuyenphoi_list(request):
    chuyenphoi = ChuyenPhoi.objects.all()
    return render(request, 'benhnhan/chuyenphoi_list.html', {'chuyenphoi': chuyenphoi})


def chuyenphoi_add(request, pk):
    data = dict()

    thongtin = get_object_or_404(ThongTin, pk=pk)
    chuyenphoi = thongtin.chuyenphoi
    data['added'] = True
    data['maso'] = thongtin.maSo
    if not chuyenphoi.added:
        chuyenphoi.ngayChuyenPhoi = date.today()
        chuyenphoi.ngayKiemTra = date.today()
        data['added'] = False

    if request.method == 'POST':
        formcp = FormCP(request.POST, instance=chuyenphoi)
        form = FormTT(request.POST, instance=thongtin)
        if formcp.is_valid():
            formcp.save()
            chuyenphoi.timeAdd = datetime.now()
            chuyenphoi.save()
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formcp = FormCP(instance=chuyenphoi)
        form = FormTT(instance=thongtin)

    context = {'formcp': formcp, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/chuyenphoi_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def chuyenphoi_edit(request, pk):
    data = dict()

    thongtin = get_object_or_404(ThongTin, pk=pk)
    chuyenphoi = thongtin.chuyenphoi
    if request.method == 'POST':
        formcp = FormCP(request.POST, instance=chuyenphoi)
        form = FormTT(instance=thongtin)
        if formcp.is_valid():
            formcp.save()
            chuyenphoi = ChuyenPhoi.objects.all()
            data['html_chuyenphoi_preview'] = render_to_string('benhnhan/includes/chuyenphoi_preview.html', {
                'chuyenphoi': chuyenphoi, 'user': request.user
            })
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formcp = FormCP(instance=chuyenphoi)
        form = FormTT(instance=thongtin)

    context = {'formcp': formcp, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/chuyenphoi_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def chuyenphoi_del(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    chuyenphoi = thongtin.chuyenphoi

    if request.method == 'POST':
        chuyenphoi.added = False
        chuyenphoi.save()
        data['form_is_valid'] = True
        chuyenphoi = ChuyenPhoi.objects.all()
        data['html_chuyenphoi_preview'] = render_to_string('benhnhan/includes/chuyenphoi_preview.html', {
            'chuyenphoi': chuyenphoi, 'user': request.user
        })
    else:
        context = {'chuyenphoi': chuyenphoi}
        data['html_form'] = render_to_string('benhnhan/includes/chuyenphoi_del.html', context, request=request)
    return JsonResponse(data)


def dongphoi_list(request):
    dongphoi = DongPhoi.objects.all()
    return render(request, 'benhnhan/dongphoi_list.html', {'dongphoi': dongphoi})


def dongphoi_add(request, pk):
    data = dict()

    thongtin = get_object_or_404(ThongTin, pk=pk)
    dongphoi = thongtin.dongphoi
    data['added'] = True
    data['maso'] = thongtin.maSo
    if not dongphoi.added:
        dtn = date.today()
        dongphoi.ngayDongPhoi = dtn
        dongphoi.ngayNopTien = dtn + relativedelta(months=6)
        data['added'] = False

    if request.method == 'POST':
        formdp = FormDP(request.POST, instance=dongphoi)
        form = FormTT(request.POST, instance=thongtin)
        if formdp.is_valid():
            formdp.save()
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formdp = FormDP(instance=dongphoi)
        form = FormTT(instance=thongtin)

    context = {'formdp': formdp, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/dongphoi_add.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def dongphoi_edit(request, pk):
    data = dict()

    thongtin = get_object_or_404(ThongTin, pk=pk)
    dongphoi = thongtin.dongphoi
    if request.method == 'POST':
        formdp = FormDP(request.POST, instance=dongphoi)
        form = FormTT(instance=thongtin)
        if formdp.is_valid():
            formdp.save()
            dongphoi = DongPhoi.objects.all()
            data['html_dongphoi_preview'] = render_to_string('benhnhan/includes/dongphoi_preview.html', {
                'dongphoi': dongphoi, 'user': request.user
            })
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        formdp = FormDP(instance=dongphoi)
        form = FormTT(instance=thongtin)

    context = {'formdp': formdp, 'form': form}
    data['html_form'] = render_to_string('benhnhan/includes/dongphoi_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)


def dongphoi_del(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    dongphoi = thongtin.dongphoi

    if request.method == 'POST':
        dongphoi.added = False
        dongphoi.save()
        data['form_is_valid'] = True
        dongphoi = DongPhoi.objects.all()
        data['html_dongphoi_preview'] = render_to_string('benhnhan/includes/dongphoi_preview.html', {
            'dongphoi': dongphoi, 'user': request.user
        })
    else:
        context = {'dongphoi': dongphoi}
        data['html_form'] = render_to_string('benhnhan/includes/dongphoi_del.html', context, request=request)
    return JsonResponse(data)


last_file_name = ''


def chochut_ex(request):
    data = dict()
    ngay = ''
    try:
        ngay = datetime.strptime(str(request.GET.get('d')), '%d/%m/%Y')
        data['chonngay'] = True
    except ValueError:
        data['chonngay'] = False

    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/landscape-template.docx')
            document._body.clear_content()
            document.add_heading(u"Ngày chọc hút: " + ngay.strftime("%d/%m/%Y"), level=1)

            chdis = ChocHut.objects.filter(gioCH__date=ngay)
            row = chdis.__len__()

            table = document.add_table(rows=row, cols=0, style='Table1')
            table.add_column(width=Cm(0.83))
            table.add_column(width=Cm(7.66))
            table.add_column(width=Cm(3.83))
            table.add_column(width=Cm(3.42))
            table.add_column(width=Cm(5.75))
            table.add_column(width=Cm(3.11))
            for i in range(row):
                ci = chdis[i]
                cell = table.cell(i, 0)
                cell.text = str(i + 1)
                cell.paragraphs[0].style = document.styles['Text2']

                cell = table.cell(i, 1)
                tv = ci.tt.tenVo.split()[-1]
                cell.text = tv
                cell.paragraphs[0].style = document.styles['Text2']
                cell.add_paragraph(ci.tt.tenVo + u' - ', style='Text4').add_run(str(ci.tt.nsVo), 't3').bold = False
                cell.add_paragraph(ci.tt.tenChong + u' - ', style='Text4').add_run(str(ci.tt.nsChong),
                                                                                   't3').bold = False

                cell = table.cell(i, 2)
                cell.text = 'HCG: ' + (ci.HCG + relativedelta(hours=7)).strftime('%H:%M')
                cell.paragraphs[0].style = document.styles['Text4']
                cell.add_paragraph('CH: ' + (ci.gioCH + relativedelta(hours=7)).strftime('%H:%M'), style='Text4')
                if ci.tt.bs is not None:
                    ten = ci.tt.bs.ten
                else:
                    ten = ''
                cell.add_paragraph('BS: ' + ten, style='Text4')
                cell.add_paragraph(str(ci.soNang) + 'N', style='Text2')

                cell = table.cell(i, 3)
                cell.text = '          M2'
                cell.paragraphs[0].style = document.styles['Text3']
                cell.add_paragraph('          M1', style='Text3')
                cell.add_paragraph('          GV', style='Text3')
                cell.add_paragraph('          TH', style='Text3')

                cell = table.cell(i, 5)
                cell.text = 'HS'
                cell.paragraphs[0].style = document.styles['Text4']
                cell.add_paragraph(ci.tt.maSo, style='Text4')

            data['form_is_valid'] = True

            document.save('docx/download.docx')
            global last_file_name
            last_file_name = 'attachment; filename="Choc-hut-' \
                             + ngay.strftime("%d-%m-%Y") \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch}
    data['html_form'] = render_to_string('benhnhan/includes/chochut_ex.html', context, request=request)
    return JsonResponse(data)


def toColor(hexx):
    if hexx == '#ff0000':
        return u'Đỏ'
    elif hexx == '#00ff00':
        return u'Xanh lá'
    elif hexx == '#0000ff':
        return u'Xanh dương'
    elif hexx == '#ffff00':
        return u'Vàng'
    elif hexx == '#fcfbe3':
        return u'Kem'
    elif hexx == '#ffffff':
        return u'Trắng'
    elif hexx == '#ffa500':
        return u'Cam'
    else:
        return u'Đã chuyển'


def trudongphoi(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    dongphoi = thongtin.dongphoi
    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/trudongphoi.docx')
            table = document.tables[0]
            ng = formch.cleaned_data['chonNgay']
            table.cell(0, 0).paragraphs[2].add_run(thongtin.maSo, 't1')
            table.cell(0, 1).paragraphs[0].add_run(u'Hà Nội, ngày ' + ng.strftime('%d') +
                                                   u' tháng ' + ng.strftime('%m') +
                                                   u' năm ' + ng.strftime('%Y'), 't2')
            table.cell(3, 0).paragraphs[0].add_run(thongtin.tenVo, 't3')
            table.cell(3, 0).paragraphs[1].add_run(thongtin.tenChong, 't3')
            table.cell(3, 0).paragraphs[2].add_run(thongtin.lienHe, 't3')
            table.cell(3, 0).paragraphs[4].add_run(dongphoi.ngayDongPhoi.strftime('%d/%m/%Y'), 't3')
            cassettle = toColor(thongtin.truphoi.label00) + ', ' + toColor(thongtin.truphoi.label01)
            table.cell(3, 0).paragraphs[6].add_run(cassettle, 't1')
            table.cell(3, 5).paragraphs[0].add_run(str(thongtin.nsVo), 't1')
            table.cell(3, 5).paragraphs[1].add_run(str(thongtin.nsChong), 't1')
            table.cell(4, 3).paragraphs[0].add_run(str(thongtin.phoi.soPhoiN2), 't3')
            table.cell(4, 3).paragraphs[1].add_run(str(thongtin.phoi.soPhoiN3), 't3')
            table.cell(4, 3).paragraphs[2].add_run(str(thongtin.phoi.soPhoiN4), 't3')
            table.cell(4, 3).paragraphs[3].add_run(str(thongtin.phoi.soPhoiN5), 't3')
            table.cell(4, 5).paragraphs[0].add_run(str(thongtin.phoi.soCryotop2), 't3')
            table.cell(4, 5).paragraphs[1].add_run(str(thongtin.phoi.soCryotop3), 't3')
            table.cell(4, 5).paragraphs[2].add_run(str(thongtin.phoi.soCryotop4), 't3')
            table.cell(4, 5).paragraphs[3].add_run(str(thongtin.phoi.soCryotop5), 't3')

            document.save('docx/download.docx')
            data['form_is_valid'] = True
            global last_file_name
            last_file_name = 'attachment; filename="Phieu-nhan-tru-dong-phoi-' \
                             + thongtin.maSo \
                             + formch.cleaned_data['chonNgay'].strftime("%d-%m-%Y") \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch, 'thongtin': thongtin}
    data['html_form'] = render_to_string('benhnhan/includes/trudongphoi.html', context, request=request)
    return JsonResponse(data)


def chuyenphoi_ex(request):
    data = dict()
    chuyenphoi = ChuyenPhoi.objects.filter(added=True)
    if request.method == 'POST':
        formcp = FormExCH(request.POST)
        if formcp.is_valid():
            document = Document('docx/cpdl.docx')
            document.add_paragraph(formcp.cleaned_data['chonNgay'].strftime("%d/%m/%Y"), style='t1')

            table = document.add_table(rows=0, cols=0, style='Table1')
            table.add_column(width=Cm(1))
            table.add_column(width=Cm(11.75))
            table.add_column(width=Cm(5.75))

            i = 1
            for cp in chuyenphoi:
                if cp.ngayRaPhoi == formcp.cleaned_data['chonNgay']:
                    row = table.add_row()
                    cell = row.cells[0]
                    cell.text = str(i)
                    cell.paragraphs[0].style = document.styles['t2']

                    cell = row.cells[1]
                    cell.text = cp.tt.tenVo.split()[-1]
                    cell.paragraphs[0].style = document.styles['t2']
                    cell.add_paragraph(cp.tt.tenVo + u' - ' + str(cp.tt.nsVo), style='t3')
                    cell.add_paragraph(cp.tt.tenChong + u' - ' + str(cp.tt.nsChong), style='t3')
                    cell.add_paragraph(' ', style='t3')
                    cell.add_paragraph(' ', style='t3')
                    if cp.tt.bs is not None:
                        ten = cp.tt.bs.ten
                    else:
                        ten = ''
                    cell.add_paragraph('BS ' + ten, style='t4')

                    cell = row.cells[2]
                    cell.text = 'HS: ' + cp.tt.maSo
                    cell.paragraphs[0].style = document.styles['t3']
                    cell.add_paragraph(' ', style='t3')
                    cell.add_paragraph(' ', style='t3')
                    cell.add_paragraph(u'RÃ ' + str(cp.soPhoiRa) + u'P', style='t4')
                    cell.add_paragraph(u'KT ' + cp.ngayKiemTra.strftime('%d/%m'), style='t4')
                    cell.add_paragraph(u'CP ' + cp.ngayChuyenPhoi.strftime('%d/%m'), style='t4')

            data['form_is_valid'] = True

            document.save('docx/download.docx')
            global last_file_name
            last_file_name = 'attachment; filename="Cpdl-' \
                             + formcp.cleaned_data['chonNgay'].strftime("%d-%m-%Y") \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formcp = FormExCH()
    context = {'formcp': formcp}
    data['html_form'] = render_to_string('benhnhan/includes/chuyenphoi_ex.html', context, request=request)
    return JsonResponse(data)


def ketquaphoi(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/ketquaphoi.docx')
            table = document.tables[0]
            table.cell(0, 4).paragraphs[0].add_run(thongtin.maSo, 't2')
            table.cell(0, 6).text = str(thongtin.chochut.stt)
            table.cell(0, 6).paragraphs[0].style = document.styles['t1']

            table.cell(2, 1).paragraphs[0].add_run(thongtin.tenVo, 't2')
            table.cell(3, 1).paragraphs[0].add_run(thongtin.tenChong, 't2')
            table.cell(2, 6).paragraphs[0].add_run(str(thongtin.nsVo), 't2')
            table.cell(3, 6).paragraphs[0].add_run(str(thongtin.nsChong), 't2')
            cell = table.cell(4, 0)
            cell.paragraphs[0].add_run(str(thongtin.trung.tongSoTrung), 't2')
            cell.paragraphs[1].add_run(str(thongtin.trung.tongSoTrungICSI), 't2')
            cell.paragraphs[2].add_run(str(thongtin.trung.tongSoTrungTT), 't2')
            cell.paragraphs[5].add_run(str(thongtin.phoi.loai1), 't2')
            cell.paragraphs[6].add_run(str(thongtin.phoi.loai2), 't2')
            cell.paragraphs[7].add_run(str(thongtin.phoi.loai3), 't2')
            cell.paragraphs[10].add_run(str(thongtin.phoi.tongSoPhoiChuyen), 't2')
            cell.paragraphs[11].add_run(str(thongtin.phoi.tongSoPhoiTiepTucTheoDoi), 't2')
            cell.paragraphs[12].add_run(str(thongtin.phoi.tongSoPhoiLuuTruLanh), 't2')
            cell.paragraphs[13].add_run(str(thongtin.phoi.tongSoPhoiHuy), 't2')
            ng = formch.cleaned_data['chonNgay']
            table.cell(6, 4).paragraphs[0].insert_paragraph_before(
                u'Ngày ' + ng.strftime('%d') +
                u' tháng ' + ng.strftime('%m') +
                u' năm ' + ng.strftime('%Y'), 't3')

            document.save('docx/download.docx')
            data['form_is_valid'] = True
            global last_file_name
            last_file_name = 'attachment; filename="Phieu-thong-bao-ve-ket-qua-phoi-' \
                             + thongtin.maSo \
                             + formch.cleaned_data['chonNgay'].strftime("%d-%m-%Y") \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch, 'thongtin': thongtin}
    data['html_form'] = render_to_string('benhnhan/includes/ketquaphoi.html', context, request=request)
    return JsonResponse(data)


def theodoiphoi(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/theodoiphoi.docx')
            table = document.tables[0]
            table.cell(0, 3).paragraphs[0].add_run(str(thongtin.maSo), 't1')
            table.cell(1, 4).paragraphs[0].add_run(str(thongtin.chochut.stt), 't3')
            table.cell(2, 0).paragraphs[0].add_run(thongtin.tenVo, 't1')
            table.cell(2, 0).paragraphs[1].add_run(thongtin.tenChong, 't1')
            table.cell(2, 2).paragraphs[0].add_run(str(thongtin.nsVo), 't1')
            table.cell(2, 2).paragraphs[1].add_run(str(thongtin.nsChong), 't1')
            table.cell(2, 3).paragraphs[0].add_run(thongtin.chochut.gioCH.strftime('%d/%m/%Y'), 't1')
            if thongtin.bs is not None:
                table.cell(2, 3).paragraphs[1].add_run(thongtin.bs.ten, 't1')
            table.cell(2, 3).paragraphs[2]\
                .add_run((thongtin.chochut.HCG + relativedelta(hours=7)).strftime('%H:%M'), 't1')
            table.cell(2, 3).paragraphs[3]\
                .add_run((thongtin.chochut.gioCH + relativedelta(hours=7)).strftime('%H:%M'), 't1')

            document.save('docx/download.docx')
            data['form_is_valid'] = True
            global last_file_name
            last_file_name = 'attachment; filename="Phieu-theo-doi-phoi-' \
                             + thongtin.maSo \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch, 'thongtin': thongtin}
    data['html_form'] = render_to_string('benhnhan/includes/theodoiphoi.html', context, request=request)
    return JsonResponse(data)


def truraphoi(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/truraphoi.docx')
            table = document.tables[0]
            table.cell(0, 10).paragraphs[0].add_run(thongtin.maSo, 't1')
            table.cell(1, 12).paragraphs[0].add_run(str(thongtin.chochut.stt), 't2')
            table.cell(3, 0).paragraphs[0].add_run(thongtin.tenVo, 't1')
            table.cell(3, 0).paragraphs[1].add_run(thongtin.tenChong, 't1')
            table.cell(3, 7).paragraphs[0].add_run(str(thongtin.nsVo), 't1')
            table.cell(3, 7).paragraphs[1].add_run(str(thongtin.nsChong), 't1')
            table.cell(3, 11).paragraphs[1].add_run(str(thongtin.phoi.tongSoPhoiLuuTruLanh), 't1')
            table.cell(4, 0).paragraphs[0].add_run(thongtin.chochut.gioCH.strftime('%d/%m/%Y'), 't1')

            document.save('docx/download.docx')
            data['form_is_valid'] = True
            global last_file_name
            last_file_name = 'attachment; filename="Phieu-theo-doi-tru-ra-phoi-' \
                             + thongtin.maSo \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch, 'thongtin': thongtin}
    data['html_form'] = render_to_string('benhnhan/includes/truraphoi.html', context, request=request)
    return JsonResponse(data)


def IVF(request, pk):
    data = dict()
    thongtin = get_object_or_404(ThongTin, pk=pk)
    if request.method == 'POST':
        formch = FormExCH(request.POST)
        if formch.is_valid():
            document = Document('docx/IVF.docx')
            table = document.tables[0]
            table.cell(0, 8).paragraphs[0].add_run(str(thongtin.chochut.stt), 't2')
            table.cell(2, 0).paragraphs[0].add_run(thongtin.tenVo, 't1')
            table.cell(2, 0).paragraphs[1].add_run(thongtin.tenChong, 't1')
            table.cell(2, 7).paragraphs[0].add_run(str(thongtin.nsVo), 't1')
            table.cell(2, 7).paragraphs[1].add_run(str(thongtin.nsChong), 't1')
            table.cell(3, 0).paragraphs[0].add_run(thongtin.chochut.gioCH.strftime('%d/%m/%Y'), 't1')
            table.cell(3, 7).paragraphs[0].add_run(thongtin.maSo, 't1')
            if thongtin.bs is not None:
                table.cell(3, 0).paragraphs[1].add_run(thongtin.bs.ten, 't1')

            document.save('docx/download.docx')
            data['form_is_valid'] = True
            global last_file_name
            last_file_name = 'attachment; filename="Bang-kiem-cac-quy-trinh-thuc-hien-IVF-' \
                             + thongtin.maSo \
                             + '.docx"'
        else:
            data['form_is_valid'] = False
    else:
        formch = FormExCH()
    context = {'formch': formch, 'thongtin': thongtin}
    data['html_form'] = render_to_string('benhnhan/includes/IVF.html', context, request=request)
    return JsonResponse(data)


def download(request):
    fs = FileSystemStorage()
    filename = 'docx/download.docx'
    global last_file_name
    with fs.open(filename) as docx:
        response = HttpResponse(docx,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = last_file_name
        return response


def change_password(request):
    data = dict()
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)  # Important!
            data['password_is_valid'] = 2
        else:
            data['password_is_valid'] = 0
    else:
        form = PasswordChangeForm(request.user)
        data['password_is_valid'] = 1
    return render(request, 'registration/change_password.html', {
        'form': form, 'data': data
    })


def thongkethang(request, thang, nam):
    thongtin = ThongTin.objects.all()
    maxy = 2000
    miny = 3000
    for tt in thongtin:
        if tt.ngayNhap.year > maxy:
            maxy = tt.ngayNhap.year
        if tt.ngayNhap.year < miny:
            miny = tt.ngayNhap.year
    years = range(miny, maxy + 1, 1)

    thongtin = ThongTin.objects.filter(ngayNhap__year=nam).filter(ngayNhap__month=thang)
    months = range(1, 13, 1)
    loai1 = 0
    loai2 = 0
    loai3 = 0
    tst = 0
    tpc = 0
    valid = True
    for tt in thongtin:
        loai1 += tt.phoi.loai1
        loai2 += tt.phoi.loai2
        loai3 += tt.phoi.loai3
        tst += tt.trung.truongThanh
        tpc += tt.phoi.tongSoPhoiChuyen
    if loai1 + loai2 + loai3 != 0:
        l1 = float(loai1) / (loai1 + loai2 + loai3)
        l2 = float(loai2) / (loai1 + loai2 + loai3)
        l3 = float(loai3) / (loai1 + loai2 + loai3)
    else:
        l1 = l2 = l3 = 1
        valid = False
    if tst == 0:
        ptst = 0
        valid = False
    else:
        ptst = float(loai1 + loai2 + loai3) / tst
    ptst *= 100

    try:
        tk = ThongKe.objects.get(nam=nam, thang=thang)
    except ThongKe.DoesNotExist:
        tk = ThongKe(nam=nam, thang=thang)
    if tk.tongSoTuiThai != 0:
        tyletresinh = tk.tongSoTreSinh/float(tk.tongSoTuiThai)*100
    else:
        tyletresinh = 0
    if tpc != 0:
        tylelamto = tk.tongSoTuiThai/float(tpc)*100
    else:
        tylelamto = 0

    return render(request, 'benhnhan/thongke.html', {'thongtin': thongtin,
                                                     'ptst': ptst,
                                                     'l1': l1,
                                                     'l2': l2,
                                                     'l3': l3,
                                                     'nam': int(nam),
                                                     'thang': int(thang),
                                                     'years': years,
                                                     'months': months,
                                                     'valid': valid,
                                                     'tyletresinh': tyletresinh,
                                                     'tylelamto': tylelamto
                                                     })


def thongkenam(request, nam):
    thongtin = ThongTin.objects.all()
    maxy = 2000
    miny = 3000
    for tt in thongtin:
        if tt.ngayNhap.year > maxy:
            maxy = tt.ngayNhap.year
        if tt.ngayNhap.year < miny:
            miny = tt.ngayNhap.year
    years = range(miny, maxy + 1, 1)

    thongtin = ThongTin.objects.filter(ngayNhap__year=nam)
    months = range(1, 13, 1)
    loai1 = [.0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0]
    loai2 = [.0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0]
    loai3 = [.0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0]
    tst = [.0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0, .0]
    l1 = 0
    l2 = 0
    l3 = 0
    tpc = [0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0]
    valid = True
    for tt in thongtin:
        k = tt.ngayNhap.month - 1
        loai1[k] += tt.phoi.loai1
        loai2[k] += tt.phoi.loai2
        loai3[k] += tt.phoi.loai3
        tst[k] += tt.trung.truongThanh
        tpc[k] += tt.phoi.tongSoPhoiChuyen

    for i in range(0, 12, 1):
        tl = loai1[i] + loai2[i] + loai3[i]
        if tl != 0:
            loai1[i] = float(loai1[i])*100/tl
            loai2[i] = float(loai2[i])*100/tl
            loai3[i] = float(loai3[i])*100/tl
            if tst[i] != 0:
                tst[i] = float(tl/tst[i])*100

    tyletresinh = [0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0]
    tylelamto = [0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0]
    for i in range(1, 13, 1):
        try:
            tk = ThongKe.objects.get(nam=nam, thang=i)
        except ThongKe.DoesNotExist:
            tk = ThongKe(nam=nam, thang=i)
            tk.save()
        if tk.tongSoTuiThai != 0:
            tyletresinh[i - 1] = tk.tongSoTreSinh/float(tk.tongSoTuiThai)*100
        else:
            tyletresinh[i - 1] = 0
        if tpc[i - 1] != 0:
            tylelamto[i - 1] = tk.tongSoTuiThai/float(tpc[i - 1])*100
        else:
            tylelamto[i - 1] = 0

    return render(request, 'benhnhan/thongkenam.html', {'thongtin': thongtin,
                                                        'tst': tst,
                                                        'l1': l1,
                                                        'l2': l2,
                                                        'l3': l3,
                                                        'loai1': loai1,
                                                        'loai2': loai2,
                                                        'loai3': loai3,
                                                        'nam': int(nam),
                                                        'years': years,
                                                        'months': months,
                                                        'valid': valid,
                                                        'tyletresinh': tyletresinh,
                                                        'tylelamto': tylelamto
                                                        })


def thongke(request):
    nam = datetime.now().year
    thang = datetime.now().month
    return redirect(str(thang) + '/' + str(nam) + '/')


def thongke_edit(request, nam):
    data = dict()
    thongtin = ThongTin.objects.all()
    maxy = 2000
    miny = 3000
    for tt in thongtin:
        if tt.ngayNhap.year > maxy:
            maxy = tt.ngayNhap.year
        if tt.ngayNhap.year < miny:
            miny = tt.ngayNhap.year

    fxo = range(0, 24, 1)

    if request.method == 'POST':
        formx = FormX(request.POST)

        if formx.is_valid():
            try:
                tk = ThongKe.objects.get(nam=nam, thang=1)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t1t1']
            tk.tongSoTuiThai = formx.cleaned_data['t1t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=2)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=2)
            tk.tongSoTreSinh = formx.cleaned_data['t2t1']
            tk.tongSoTuiThai = formx.cleaned_data['t2t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=3)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=3)
            tk.tongSoTreSinh = formx.cleaned_data['t3t1']
            tk.tongSoTuiThai = formx.cleaned_data['t3t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=4)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=4)
            tk.tongSoTreSinh = formx.cleaned_data['t4t1']
            tk.tongSoTuiThai = formx.cleaned_data['t4t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=5)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=5)
            tk.tongSoTreSinh = formx.cleaned_data['t5t1']
            tk.tongSoTuiThai = formx.cleaned_data['t5t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=6)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t6t1']
            tk.tongSoTuiThai = formx.cleaned_data['t6t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=7)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t7t1']
            tk.tongSoTuiThai = formx.cleaned_data['t7t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=8)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t8t1']
            tk.tongSoTuiThai = formx.cleaned_data['t8t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=9)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t9t1']
            tk.tongSoTuiThai = formx.cleaned_data['t9t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=10)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t10t1']
            tk.tongSoTuiThai = formx.cleaned_data['t10t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=11)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t11t1']
            tk.tongSoTuiThai = formx.cleaned_data['t11t2']
            tk.save()
            try:
                tk = ThongKe.objects.get(nam=nam, thang=12)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=1)
            tk.tongSoTreSinh = formx.cleaned_data['t12t1']
            tk.tongSoTuiThai = formx.cleaned_data['t12t2']
            tk.save()

            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
    else:
        for i in range(0, 12, 1):
            try:
                tk = ThongKe.objects.get(nam=nam, thang=i + 1)
            except ThongKe.DoesNotExist:
                tk = ThongKe(nam=nam, thang=i + 1)
                tk.save()
            fxo[i*2] = tk.tongSoTreSinh
            fxo[i*2 + 1] = tk.tongSoTuiThai

        formx = FormX(initial={'t1t1': fxo[0], 't1t2': fxo[1],
                               't2t1': fxo[2], 't2t2': fxo[3],
                               't3t1': fxo[4], 't3t2': fxo[5],
                               't4t1': fxo[6], 't4t2': fxo[7],
                               't5t1': fxo[8], 't5t2': fxo[9],
                               't6t1': fxo[10], 't6t2': fxo[11],
                               't7t1': fxo[12], 't7t2': fxo[13],
                               't8t1': fxo[14], 't8t2': fxo[15],
                               't9t1': fxo[16], 't9t2': fxo[17],
                               't10t1': fxo[18], 't10t2': fxo[19],
                               't11t1': fxo[20], 't11t2': fxo[21],
                               't12t1': fxo[22], 't12t2': fxo[23]
                               })

    context = {'formx': formx, 'nam': nam}
    data['html_form'] = render_to_string('benhnhan/includes/thongke_edit.html',
                                         context,
                                         request=request
                                         )
    return JsonResponse(data)
